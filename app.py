import os
import sys
import logging
import tempfile
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import DocArrayInMemorySearch
from langchain.embeddings import OpenAIEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.docstore.document import Document
from docx import Document as DocxDocument
import pandas as pd
import time
from pinecone import Pinecone as PineconeClient
from langchain_pinecone import PineconeVectorStore
import re
import requests

# ---------------------------------------
# 1. Loglama Ayarları
# ---------------------------------------
if not os.path.exists("logs"):
    os.makedirs("logs")

logging.basicConfig(
    filename="logs/app.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Hata loglaması için özel handler
error_handler = logging.FileHandler("logs/error.log")
error_handler.setLevel(logging.ERROR)
error_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
logging.getLogger().addHandler(error_handler)

# ---------------------------------------
# 2. Streamlit Sayfa Yapısı
# ---------------------------------------
st.set_page_config(page_title="Villa Villa Yapay Zeka", layout="wide", 
                   initial_sidebar_state="expanded")

# CSS ile özelleştirme
st.markdown("""
<style>
    .main-header {
        display: flex;
        align-items: center;
        background-color: #f5f5f5;
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .chat-message-user {
        background-color: #e6f7ff;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .chat-message-assistant {
        background-color: #f0f0f0;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------
# 3. API Anahtarları Yönetimi
# ---------------------------------------
def load_api_keys():
    # OpenAI API anahtarını al
    openai_api_key = st.secrets.get("openai", {}).get("api_key", None)
    if not openai_api_key:
        st.error("OpenAI API anahtarı bulunamadı. Lütfen Streamlit Secrets ayarlarını kontrol edin.")
        logging.error("OpenAI API anahtarı bulunamadı")
        return False
    
    # Pinecone API anahtarını al
    pinecone_api_key = st.secrets.get("pinecone", {}).get("api_key", None)
    pinecone_environment = st.secrets.get("pinecone", {}).get("environment", None)
    pinecone_index_name = st.secrets.get("pinecone", {}).get("index_name", None)
    
    if not pinecone_api_key or not pinecone_environment or not pinecone_index_name:
        st.error("Pinecone API bilgileri eksik. Lütfen Streamlit Secrets ayarlarını kontrol edin.")
        logging.error("Pinecone API bilgileri eksik")
        return False
    
    os.environ["OPENAI_API_KEY"] = openai_api_key
    os.environ["PINECONE_API_KEY"] = pinecone_api_key
    os.environ["PINECONE_ENVIRONMENT"] = pinecone_environment
    os.environ["PINECONE_INDEX_NAME"] = pinecone_index_name
    
    return True

# ---------------------------------------
# 4. Google Drive Doküman İndirme
# ---------------------------------------
def extract_document_id(url):
    """Google Doküman URL'sinden doküman ID'sini çıkarır."""
    pattern = r"/d/([a-zA-Z0-9-_]+)"
    match = re.search(pattern, url)
    if match:
        return match.group(1)
    return None

def download_google_doc(doc_url, file_format="docx"):
    """Google Dokümanı doğrudan export formatında indirir."""
    try:
        doc_id = extract_document_id(doc_url)
        if not doc_id:
            logging.error(f"Geçersiz Google Doküman URL'si: {doc_url}")
            return None, None
        
        # Google Docs'un export URL'si
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format={file_format}"
        
        # Dokümanı indir
        response = requests.get(export_url)
        
        # HTTP yanıt kodunu ve içerik uzunluğunu logla
        logging.info(f"Doküman indirme yanıtı: {response.status_code}, İçerik uzunluğu: {len(response.content)}")
        
        if response.status_code != 200:
            logging.error(f"Doküman indirilemedi. Durum kodu: {response.status_code}")
            return None, None
        
        # Geçici dosya oluştur
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_format}")
        temp_file.write(response.content)
        temp_file.close()
        
        return temp_file.name, doc_id
        
    except Exception as e:
        logging.error(f"Doküman indirme hatası: {str(e)}")
        return None, None

def process_document(file_path, doc_name, file_format="docx"):
    """İndirilen dokümanı işleyerek LangChain Document nesnesine dönüştürür."""
    try:
        if file_format == "docx":
            # DOCX dosyasını oku
            docx = DocxDocument(file_path)
            text = "\n".join([p.text for p in docx.paragraphs if p.text.strip() != ""])
            
            # Tablo içeriklerini de ekle
            for table in docx.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += f"\n{row_text}"
            
            return Document(page_content=text, metadata={"source": doc_name})
            
        elif file_format == "txt":
            # TXT dosyasını oku
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return Document(page_content=text, metadata={"source": doc_name})
            
        else:
            logging.error(f"Desteklenmeyen dosya formatı: {file_format}")
            return None
            
    except Exception as e:
        logging.error(f"Doküman işleme hatası: {str(e)}")
        return None
    finally:
        # Geçici dosyayı temizle
        try:
            os.remove(file_path)
        except:
            pass

def load_documents_from_urls(doc_urls):
    """Verilen URL listesinden Google Dokümanları yükler."""
    documents = []

    # Yükleme durumu göstergesi
    progress_bar = st.progress(0)
    st.write(f"Google Drive'dan {len(doc_urls)} doküman yükleniyor...")

    # Başarılı ve başarısız dokümanları izle
    successful_docs = []
    failed_docs = []

    for idx, (doc_name, doc_url) in enumerate(doc_urls.items()):
        try:
            # İlk olarak DOCX formatında indirmeyi dene
            file_path, doc_id = download_google_doc(doc_url, "docx")

            # DOCX indirilemediyse, TXT formatında dene
            if not file_path:
                st.warning(f"{doc_name} DOCX olarak indirilemedi, TXT olarak deneniyor...")
                file_path, doc_id = download_google_doc(doc_url, "txt")

                if not file_path:
                    failed_docs.append(doc_name)
                    st.error(f"{doc_name} dokümanı indirilemedi! Dokümanın paylaşım ayarlarını kontrol edin.")
                    continue

                # TXT dosyasını işle
                document = process_document(file_path, doc_name, "txt")
            else:
                # DOCX dosyasını işle
                document = process_document(file_path, doc_name, "docx")

            if document:
                documents.append(document)
                successful_docs.append(doc_name)
                logging.info(f"Doküman başarıyla yüklendi: {doc_name}, İçerik uzunluğu: {len(document.page_content)} karakter")
            else:
                failed_docs.append(doc_name)
                st.error(f"{doc_name} dokümanı işlenemedi!")

            # İlerleme durumunu güncelle
            progress_bar.progress((idx + 1) / len(doc_urls))

        except Exception as e:
            failed_docs.append(doc_name)
            logging.error(f"{doc_name} işlenirken hata: {str(e)}")
            st.error(f"{doc_name} için beklenmeyen bir hata oluştu.")

    # Sonuçları göster
    if successful_docs:
        st.success(f"Toplam {len(successful_docs)} belge başarıyla yüklendi: {', '.join(successful_docs)}")

    if failed_docs:
        st.error(f"Toplam {len(failed_docs)} belge yüklenemedi: {', '.join(failed_docs)}")

    return documents

# ---------------------------------------
# 5. Pinecone Vektör Veritabanı İşlemleri
# ---------------------------------------
def init_pinecone():
    """Pinecone'u başlatır ve varsa indeksi döndürür."""
    try:
        # Pinecone API bilgilerini al
        api_key = os.environ.get("PINECONE_API_KEY")
        environment = os.environ.get("PINECONE_ENVIRONMENT")
        index_name = os.environ.get("PINECONE_INDEX_NAME")
        
        if not api_key or not environment or not index_name:
            st.error("Pinecone API bilgileri eksik.")
            return None, None
        
        # Pinecone'u başlat
        pc = PineconeClient(api_key=api_key)
        
        # İndeksi kontrol et
        indexes = pc.list_indexes()
        if index_name not in indexes:
            st.info(f"Pinecone indeksi '{index_name}' bulunamadı, oluşturuluyor...")
            # OpenAI embedding modeli için 1536 boyut kullanılır
            pc.create_index(
                name=index_name,
                dimension=3072,  # text-embedding-3-large için 3072
                metric="cosine"
            )
            st.success(f"Pinecone indeksi '{index_name}' başarıyla oluşturuldu.")
            
        # İndekse bağlan
        index = pc.Index(index_name)
        
        return index, index_name
        
    except Exception as e:
        st.error(f"Pinecone başlatma hatası: {str(e)}")
        logging.error(f"Pinecone başlatma hatası: {str(e)}")
        return None, None

def create_or_update_vector_db(documents, namespace="villa_villa"):
    """Belgeleri Pinecone'a ekler veya günceller."""
    try:
        # OpenAI embeddings oluştur
        embeddings = OpenAIEmbeddings(
            model="text-embedding-3-large"  # 3072 boyutlu embedding oluşturmak için
        )
        
        # Pinecone'u başlat
        index, index_name = init_pinecone()
        if not index or not index_name:
            return None
        
        # Belgeleri parçalara böl
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=2500, 
            chunk_overlap=400,
            separators=["\n\n", "\n", ". ", " ", ""],
            is_separator_regex=False,
            length_function=len
        )
        chunks = splitter.split_documents(documents)
        logging.info(f"Belgeler {len(chunks)} parçaya bölündü")
        
        # Namespace'teki mevcut vektörleri temizle
        try:
            stats = index.describe_index_stats()
            if namespace in stats.get("namespaces", {}):
                st.info(f"Pinecone namespace '{namespace}' temizleniyor...")
                # Mevcut namespace vektörlerini temizle
                vector_count = stats["namespaces"][namespace]["vector_count"]
                
                if vector_count > 0:
                    # Boş bir sorgu yaparak tüm verileri al ve ID'leri topla
                    dummy_vector = [0.0] * 3072  # text-embedding-3-large için 3072 boyut
                    results = index.query(
                        vector=dummy_vector,
                        namespace=namespace,
                        top_k=min(vector_count, 10000),
                        include_metadata=False,
                        include_values=False
                    )
                    
                    # ID'leri topla ve sil
                    if hasattr(results, 'matches') and results.matches:
                        id_list = [match.id for match in results.matches]
                        if id_list:
                            index.delete(ids=id_list, namespace=namespace)
                            st.success(f"Namespace '{namespace}' temizlendi, {len(id_list)} vektör silindi.")
        except Exception as e:
            st.warning(f"Namespace temizleme sırasında hata: {str(e)}. İşlem devam ediyor...")
            logging.error(f"Namespace temizleme hatası: {str(e)}")
        
        # Vektör veritabanı oluştur
        with st.spinner(f"Vektörler Pinecone'a yükleniyor ({len(chunks)} parça)..."):
            text_field = "text"  # Metin alanının adı
            
            vectorstore = PineconeVectorStore(
                index=index,
                embedding=embeddings,
                text_field=text_field,
                namespace=namespace
            )
            
            # Parçaları ekle
            vectorstore.add_documents(chunks)
            
            st.success(f"Vektörler Pinecone'a başarıyla yüklendi!")
            
            return vectorstore
            
    except Exception as e:
        st.error(f"Vektör veritabanı oluşturma hatası: {str(e)}")
        logging.error(f"Vektör veritabanı oluşturma hatası: {str(e)}")
        return None

def load_vector_db_from_pinecone(namespace="villa_villa"):
    """Pinecone'dan vektör veritabanını yükler."""
    try:
        # OpenAI embeddings oluştur
        embeddings = OpenAIEmbeddings(
            model="text-embedding-3-large"  # 3072 boyutlu embedding
        )
        
        # Pinecone'u başlat
        index, index_name = init_pinecone()
        if not index or not index_name:
            return None
        
        # İndeks istatistiklerini kontrol et
        stats = index.describe_index_stats()
        if namespace not in stats.get("namespaces", {}) or stats["namespaces"][namespace]["vector_count"] == 0:
            st.warning(f"Pinecone'da '{namespace}' namespace'inde veri bulunamadı. Lütfen önce verileri yükleyin.")
            return None
        
        # Pinecone vektör deposunu yükle
        vectorstore = PineconeVectorStore(
            index=index,
            embedding=embeddings,
            text_field="text",
            namespace=namespace
        )
        
        vector_count = stats["namespaces"][namespace]["vector_count"]
        st.success(f"Pinecone vektör veritabanı başarıyla yüklendi! Toplam {vector_count} vektör bulundu.")
        
        return vectorstore
        
    except Exception as e:
        st.error(f"Pinecone'dan vektör veritabanı yükleme hatası: {str(e)}")
        logging.error(f"Pinecone'dan vektör veritabanı yükleme hatası: {str(e)}")
        return None

# ---------------------------------------
# 6. Özel Prompt Şablonu
# ---------------------------------------
def create_qa_prompt():
    template = """
    Sen Villa Villa şirketinin finans ve operasyon asistanısın. Verilen bilgilere dayanarak kapsamlı ve doğru yanıtlar vermelisin.

    # Görevin
    - İşletme finansları, tedarikçi bilgileri, giderler, personel bilgileri ve geçmiş işlerle ilgili sorulara cevap ver
    - Tüm cevaplarını aşağıdaki belgelerden aldığın bilgilere dayandır
    - Emin olmadığın konularda tahmin yürütme
    - Bilginin eksik olduğu durumlarda dürüstçe bilmediğini söyle

    # Belgelerden Alınan Bilgiler:
    {context}

    # Sohbet Geçmişi:
    {chat_history}

    # Mevcut Soru:
    {question}

    # Yanıt Kuralları:
    1. Bilgiyi net ve düzgün bir şekilde organize et
    2. Sayısal veriler için tablolar kullanmayı tercih et
    3. Detaylı finansal analizlerde bilgileri kategorilere ayır
    4. Tahmini veriler sunacaksan, bunun tahmini olduğunu açıkça belirt
    5. Bilgi yoksa "Bu konuda mevcut belgelerimde bilgi bulamadım. Villa Villa yönetimine danışmanızı öneririm." de

    # Yanıtını aşağıya yaz:
    """
    return PromptTemplate(input_variables=["context", "chat_history", "question"], template=template)

# ---------------------------------------
# 7. Chat Zinciri Kurulumu
# ---------------------------------------
def create_chat_chain(vector_db):
    try:
        llm = ChatOpenAI(
            temperature=0.2,
            model_name="gpt-4-turbo",
            verbose=True,
            streaming=True
        )
        
        # Retriever - Similarity arama ve daha fazla belge getirme
        retriever = vector_db.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 20}
        )
        
        qa_prompt = create_qa_prompt()
        
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            return_source_documents=True,
            combine_docs_chain_kwargs={"prompt": qa_prompt},
            verbose=True
        )
        return chain
    except Exception as e:
        logging.error(f"Chat zinciri hatası: {str(e)}")
        return None

# ---------------------------------------
# 8. Ana Uygulama
# ---------------------------------------
def main():
    # Session state değişkenleri
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    
    # Başlık bölümü
    with st.container():
        st.markdown('<div class="main-header">', unsafe_allow_html=True)
        col1, col2 = st.columns([1, 5])
        with col1:
            try:
                st.image("assets/villa_villa_logo.jpg", width=100)
            except Exception as e:
                logging.error(f"Logo yüklenirken hata: {str(e)}")
        with col2:
            st.title("Villa Villa Yapay Zeka ile Sohbet")
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Sidebar bilgileri
    with st.sidebar:
        st.subheader("Villa Villa Asistan Hakkında")
        st.info("""Bu yapay zeka asistanı, Villa Villa şirketinin finansal ve 
        operasyonel bilgilerine dayanarak sorularınızı yanıtlamak için tasarlanmıştır.
        Tedarikçi bilgileri, gider analizleri, personel bilgileri ve daha fazlası 
        hakkında sorular sorabilirsiniz.""")
        
        st.subheader("Örnek Sorular")
        st.markdown("""
        - Kasım 2024'teki toplam giderler nelerdir?
        - En büyük tedarikçimiz hangisidir?
        - Personel maaşları ne kadardır?
        - Araç giderleri nelerdir?
        - Metro'dan yapılan alışverişlerin toplam tutarı nedir?
        """)
        
        # Veri yükleme ayarları
        st.subheader("Veri Yönetimi")
        load_option = st.radio(
            "Veri kaynağını seçin:",
            ["Pinecone'dan Yükle", "Google Drive'dan Yükle ve Pinecone'a Kaydet"]
        )
        
        if load_option == "Google Drive'dan Yükle ve Pinecone'a Kaydet":
            custom_namespace = st.text_input(
                "Pinecone namespace (isteğe bağlı)",
                value="villa_villa",
                help="Vektörlerin kaydedileceği namespace adı"
            )
        else:
            custom_namespace = st.text_input(
                "Pinecone namespace (isteğe bağlı)",
                value="villa_villa",
                help="Yüklenecek vektörlerin bulunduğu namespace adı"
            )
        
        # Veri yenileme butonu
        refresh_data = st.button("🔄 Verileri Yenile", use_container_width=True)
    
    # API anahtarlarını yükle
    if not load_api_keys():
        st.stop()
    
    # Google Drive doküman URL'leri
    doc_urls = {
        "gelen_faturalar": "https://docs.google.com/document/d/1TfyGyepmojRdD6xd7WD73I8BEKLypwtzGOa9tvVVtPE/edit",
        "genel_gider": "https://docs.google.com/document/d/1TkQsG2f9BBIiSiIE_sI-Q2k9cNoD8PzxcSf1qq-izow/edit",
        "personel_giderleri": "https://docs.google.com/document/d/1F9xxY5VztoBi7lqH95jQ-TzOTBGH00_5y-ZbcHGZTHI/edit",
        "villa_villa_tanitim": "https://docs.google.com/document/d/16rXwlBEkjbH2pEcUgtseuYMhvZONZkitWGhgVtNDJsY/edit",
        "yapilan_isler": "https://docs.google.com/document/d/1D6jDry4yEeWEWpluDMqOTmqNuBc449Oc84hcVIEqf1w/edit"
    }
    
    # Veriler yeniden yüklensin mi?
    if refresh_data or "vector_db" not in st.session_state:
        try:
            if load_option == "Google Drive'dan Yükle ve Pinecone'a Kaydet":
                # Google Drive'dan belgeleri yükle
                with st.spinner("Google Drive'dan belgeler yükleniyor..."):
                    documents = load_documents_from_urls(doc_urls)
                    if not documents:
                        st.error("Hiçbir doküman yüklenemedi! URL'leri ve dokümanların paylaşım ayarlarını kontrol edin.")
                        st.stop()
                    
                    logging.info(f"Toplam {len(documents)} belge Drive'dan yüklendi")
                
                # Vektör veritabanı oluştur ve Pinecone'a kaydet
                with st.spinner("Vektör veritabanı oluşturuluyor ve Pinecone'a kaydediliyor..."):
                    vector_db = create_or_update_vector_db(documents, namespace=custom_namespace)
                    if not vector_db:
                        st.error("Vektör veritabanı oluşturulamadı!")
                        st.stop()
                    
                    st.session_state.vector_db = vector_db
            else:
                # Pinecone'dan var olan veritabanını yükle
                with st.spinner("Pinecone'dan vektör veritabanı yükleniyor..."):
                    vector_db = load_vector_db_from_pinecone(namespace=custom_namespace)
                    if not vector_db:
                        st.error("Pinecone'dan vektör veritabanı yüklenemedi! Önce verileri yüklediğinizden emin olun.")
                        st.stop()
                    
                    st.session_state.vector_db = vector_db
            
            # Chat zinciri oluştur
            with st.spinner("Sohbet sistemi hazırlanıyor..."):
                chat_chain = create_chat_chain(st.session_state.vector_db)
                if not chat_chain:
                    st.error("Sohbet sistemi oluşturulamadı!")
                    st.stop()
                
                st.session_state.chat_chain = chat_chain
                
            st.success("Sistem başarıyla yüklendi! Sorunuzu sorabilirsiniz.")
            time.sleep(1)
            st.rerun()
            
        except Exception as e:
            st.error(f"Uygulama başlatılırken bir hata oluştu: {str(e)}")
            logging.error(f"Uygulama başlatılırken bir hata oluştu: {str(e)}")
            st.stop()
    
    # Sohbet geçmişini görüntüle
    chat_container = st.container()
    with chat_container:
        for i in range(0, len(st.session_state.chat_history), 2):
            if i < len(st.session_state.chat_history):
                with st.chat_message("user", avatar="👤"):
                    st.markdown(st.session_state.chat_history[i][1])
            
            if i+1 < len(st.session_state.chat_history):
                with st.chat_message("assistant", avatar="🏛️"):
                    st.markdown(st.session_state.chat_history[i+1][1])
    
# Kullanıcı girişi
user_input = st.chat_input("Villa Villa hakkında bir soru sorun...")

# Temizleme butonları
cols = st.columns(2)
with cols[0]:
    if st.button("🧹 Sohbeti Temizle", use_container_width=True):
        st.session_state.chat_history = []
        st.rerun()
with cols[1]:
    if st.button("🔄 Önbelleği Yenile", use_container_width=True):
        chat_history = st.session_state.chat_history
        for key in list(st.session_state.keys()):
            if key != "chat_history":
                del st.session_state[key]
        st.session_state.chat_history = chat_history
        st.rerun()

if user_input:
    logging.info(f"Kullanıcı sorusu: {user_input}")
    
    with st.chat_message("user", avatar="👤"):
        st.markdown(user_input)
    
    st.session_state.chat_history.append(("user", user_input))

    try:
        # Sohbet geçmişini uygun formata dönüştür
        chat_formatted = []
        for i in range(0, len(st.session_state.chat_history)-1, 2):
            if i+1 < len(st.session_state.chat_history):
                chat_formatted.append((st.session_state.chat_history[i][1],
                                       st.session_state.chat_history[i+1][1]))

        # Yanıt oluştur
        message_placeholder = st.empty()
        with st.chat_message("assistant", avatar="🏛️"):
            message_placeholder = st.empty()
            full_response = ""

            with st.spinner("Villa Villa Asistanı düşünüyor..."):
                response = st.session_state.chat_chain({
                    "question": user_input,
                    "chat_history": chat_formatted
                })
                full_response = response["answer"]

                # Kaynakları logla
                if "source_documents" in response:
                    sources = [doc.metadata.get("source", "Bilinmeyen Kaynak")
                               for doc in response["source_documents"]]
                    logging.info(f"Yanıt kaynakları: {set(sources)}")

                # Yanıtı göster
                message_placeholder.markdown(full_response)

        st.session_state.chat_history.append(("assistant", full_response))

    except Exception as e:
        error_msg = f"Yanıt oluşturma hatası: {str(e)}"
        logging.error(error_msg)
        with st.chat_message("assistant", avatar="🏛️"):
            st.error("Üzgünüm, yanıt oluşturulurken bir hata oluştu.")
        st.session_state.chat_history.append(("assistant", "Üzgünüm, bir hata oluştu."))

# ---------------------------------------
# 9. Uygulama Başlatılıyor
# ---------------------------------------
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logging.critical(f"Kritik uygulama hatası: {str(e)}")
        try:
            st.error("Beklenmeyen bir hata oluştu. Lütfen logs klasörünü kontrol edin veya uygulamayı yeniden başlatın.")
        except:
            print("Streamlit dışında kritik hata:", e)

