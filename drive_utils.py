import re
import os
import tempfile
import requests
import logging
from docx import Document as DocxDocument
from langchain.docstore.document import Document
import streamlit as st

def extract_document_id(url):
    """Google Doküman URL'sinden doküman ID'sini çıkarır."""
    pattern = r"/d/([a-zA-Z0-9-_]+)"
    match = re.search(pattern, url)
    return match.group(1) if match else None

def download_google_doc(doc_url, file_format="docx"):
    """Google Dokümanı belirtilen formatta indirir (varsayılan: DOCX)."""
    try:
        doc_id = extract_document_id(doc_url)
        if not doc_id:
            logging.error(f"Geçersiz Google Doküman URL'si: {doc_url}")
            return None, None

        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format={file_format}"
        response = requests.get(export_url)

        if response.status_code != 200:
            logging.error(f"Doküman indirilemedi: {response.status_code}")
            return None, None

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_format}")
        temp_file.write(response.content)
        temp_file.close()

        return temp_file.name, doc_id

    except Exception as e:
        logging.exception("Doküman indirme hatası")
        return None, None

def process_document(file_path, doc_name, file_format="docx"):
    """İndirilen dosyayı LangChain Document formatına dönüştürür."""
    try:
        if file_format == "docx":
            docx = DocxDocument(file_path)
            text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
            for table in docx.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += f"\n{row_text}"
        elif file_format == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            logging.error(f"Desteklenmeyen format: {file_format}")
            return None

        return Document(page_content=text, metadata={"source": doc_name})

    except Exception as e:
        logging.exception("Doküman işleme hatası")
        return None

    finally:
        try:
            os.remove(file_path)
        except:
            pass

def load_documents_from_urls(doc_urls):
    """Birden fazla Google Docs bağlantısından belge yükler."""
    documents = []
    successful_docs = []
    failed_docs = []

    progress_bar = st.progress(0)
    st.write(f"{len(doc_urls)} doküman indiriliyor...")

    for idx, (doc_name, doc_url) in enumerate(doc_urls.items()):
        try:
            file_path, _ = download_google_doc(doc_url, "docx")
            if not file_path:
                st.warning(f"{doc_name} DOCX indirilemedi, TXT deneniyor...")
                file_path, _ = download_google_doc(doc_url, "txt")
                if not file_path:
                    st.error(f"{doc_name} indirilemedi. Paylaşım ayarlarını kontrol edin.")
                    failed_docs.append(doc_name)
                    continue
                document = process_document(file_path, doc_name, "txt")
            else:
                document = process_document(file_path, doc_name, "docx")

            if document:
                documents.append(document)
                successful_docs.append(doc_name)
                logging.info(f"{doc_name} yüklendi ({len(document.page_content)} karakter)")
            else:
                st.error(f"{doc_name} işlenemedi.")
                failed_docs.append(doc_name)

            progress_bar.progress((idx + 1) / len(doc_urls))

        except Exception as e:
            logging.exception(f"{doc_name} için hata oluştu.")
            st.error(f"{doc_name} yüklenemedi.")
            failed_docs.append(doc_name)

    if successful_docs:
        st.success(f"{len(successful_docs)} doküman yüklendi: {', '.join(successful_docs)}")
    if failed_docs:
        st.error(f"{len(failed_docs)} doküman başarısız: {', '.join(failed_docs)}")

    return documents
